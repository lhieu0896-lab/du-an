# -*- coding: utf-8 -*-
import sys
import os
import io
import base64
import pandas as pd
import docx
import bcrypt
from pypdf import PdfReader
from groq import Groq
from supabase import create_client, Client
from duckduckgo_search import DDGS
import streamlit as st

# Xuất utf-8 tránh lỗi font chữ tiếng Việt
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

class DongCoAI:
    def __init__(self):
        # Lấy danh sách API Keys từ secrets (Ép kiểu chuẩn chống gạch vàng VS Code)
        self.danh_sach_keys = []
        try:
            if "GROQ_API_KEYS" in st.secrets:
                cac_keys = st.secrets["GROQ_API_KEYS"]
                if isinstance(cac_keys, (list, tuple)):
                    self.danh_sach_keys = [str(k).strip() for k in cac_keys if str(k).strip()]
                else:
                    self.danh_sach_keys = [str(cac_keys).strip()]
            elif "GROQ_API_KEY" in st.secrets:
                key_don = str(st.secrets["GROQ_API_KEY"]).strip()
                if key_don:
                    self.danh_sach_keys = [key_don]
        except Exception:
            pass

        if not self.danh_sach_keys:
            key_don = os.getenv("GROQ_API_KEY", "")
            if key_don:
                self.danh_sach_keys = [key_don]

        # Kết nối CSDL Supabase
        self.ket_noi_supabase: Client = None
        self.trang_thai_db = "CHƯA_KẾT_NỐI"
        try:
            url_db = st.secrets.get("SUPABASE_URL") or os.getenv("SUPABASE_URL", "")
            khoa_db = st.secrets.get("SUPABASE_KEY") or os.getenv("SUPABASE_KEY", "")

            if url_db and khoa_db:
                self.ket_noi_supabase = create_client(url_db, khoa_db)
                self.trang_thai_db = "ĐÃ_KẾT_NỐI"
        except Exception as e:
            self.trang_thai_db = f"LỖI: {str(e)}"

        # Chuyển mô hình văn bản sang llama-3.1-8b-instant (tốc độ cao & hạn mức rộng)
        self.mo_hinh_van_ban = "llama-3.1-8b-instant"
        self.mo_hinh_hinh_anh = "llama-3.2-11b-vision-preview"

    # --- ĐĂNG KÝ / ĐĂNG NHẬP ---
    def dang_ky_tai_khoan(self, ho_ten, ten_dang_nhap, mat_khau):
        if not self.ket_noi_supabase:
            return False, "Chưa kết nối CSDL!"
        try:
            ten_dang_nhap_chuandanh = ten_dang_nhap.strip().lower()
            ket_qua = self.ket_noi_supabase.table("users").select("*").eq("username", ten_dang_nhap_chuandanh).execute()
            if ket_qua.data:
                return False, "Tài khoản này đã tồn tại!"
            
            mat_khau_ma_hoa = bcrypt.hashpw(mat_khau.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            
            self.ket_noi_supabase.table("users").insert({
                "full_name": ho_ten.strip(),
                "username": ten_dang_nhap_chuandanh,
                "password_hash": mat_khau_ma_hoa
            }).execute()
            
            return True, "Đăng ký thành công!"
        except Exception as e:
            return False, f"Lỗi đăng ký: {str(e)}"

    def xac_thuc_dang_nhap(self, ten_dang_nhap, mat_khau):
        if not self.ket_noi_supabase:
            return None, "Chưa kết nối CSDL!"
        try:
            ten_dang_nhap_chuandanh = ten_dang_nhap.strip().lower()
            ket_qua = self.ket_noi_supabase.table("users").select("*").eq("username", ten_dang_nhap_chuandanh).execute()
            if not ket_qua.data:
                return None, "Sai tên tài khoản hoặc mật khẩu!"
            
            thong_tin_nguoi_dung = ket_qua.data[0]
            mat_khau_da_luu = thong_tin_nguoi_dung["password_hash"].encode('utf-8')
            
            if bcrypt.checkpw(mat_khau.encode('utf-8'), mat_khau_da_luu):
                return thong_tin_nguoi_dung, "Đăng nhập thành công!"
            else:
                return None, "Sai tên tài khoản hoặc mật khẩu!"
        except Exception as e:
            return None, f"Lỗi đăng nhập: {str(e)}"

    # --- TÌM KIẾM WEB REAL-TIME ---
    def tim_kiem_web(self, cau_hoi, so_ket_qua_toi_da=5):
        try:
            danh_sach_ket_qua = []
            with DDGS() as tim_kiem:
                ket_qua_raw = list(tim_kiem.text(cau_hoi, region="vn-vi", max_results=so_ket_qua_toi_da))
                if not ket_qua_raw:
                    ket_qua_raw = list(tim_kiem.text(cau_hoi, max_results=so_ket_qua_toi_da))
                
                for r in ket_qua_raw:
                    tieu_de = r.get('title', '')
                    duong_dan = r.get('href', '')
                    noi_dung = r.get('body', '')
                    if noi_dung:
                        danh_sach_ket_qua.append(f"📌 **{tieu_de}**\nNguồn: {duong_dan}\nNội dung: {noi_dung}\n")
                        
            return "\n".join(danh_sach_ket_qua) if danh_sach_ket_qua else "CHƯA_CÓ_KẾT_QUẢ"
        except Exception as e:
            return f"Lỗi tìm kiếm web: {str(e)}"

    # --- TỰ ĐỘNG ĐẶT TÊN CHAT (AN TOÀN CHỐNG NGHẼN) ---
    def tao_tieu_de_chat(self, cau_hoi_dau_tien):
        cat_ngan = cau_hoi_dau_tien.strip()[:20]
        return f"Chat: {cat_ngan}..." if cat_ngan else "Cuộc trò chuyện mới"

    # --- BỘ LỌC TÀI LIỆU DÀI (LIGHT RAG) ---
    def trich_xuat_doan_lien_quan(self, van_ban_goc, cau_hoi, kich_thuoc_doan=1500, so_doan_lay=3):
        if len(van_ban_goc) <= kich_thuoc_doan * 2:
            return van_ban_goc

        danh_sach_doan = [van_ban_goc[i:i+kich_thuoc_doan] for i in range(0, len(van_ban_goc), kich_thuoc_doan - 200)]
        tu_khoa_cau_hoi = set(cau_hoi.lower().split())
        doan_co_diem = []
        
        for doan in danh_sach_doan:
            doan_chuyen_thuong = doan.lower()
            diem = sum(1 for tu in tu_khoa_cau_hoi if len(tu) > 2 and tu in doan_chuyen_thuong)
            doan_co_diem.append((diem, doan))
        
        doan_co_diem.sort(key=lambda x: x[0], reverse=True)
        doan_duoc_chon = [muc[1] for muc in doan_co_diem[:so_doan_lay]]
        return "\n\n...[Đoạn trích xuất từ tài liệu]...\n\n".join(doan_duoc_chon)

    # --- XUẤT CHAT RA FILE WORD ---
    def xuat_chat_ra_file_word(self, danh_sach_tin_nhan, tieu_de_chat):
        tai_lieu = docx.Document()
        tai_lieu.add_heading(f"BÁO CÁO TRÒ CHUYỆN: {tieu_de_chat}", level=1)
        tai_lieu.add_paragraph("Trích xuất từ Hệ thống NEXUS AI\n-----------------------------------")
        
        for tin_nhan in danh_sach_tin_nhan:
            ten_vai_tro = "NGƯỜI DÙNG" if tin_nhan["role"] == "user" else "TRỢ LÝ AI"
            noi_dung = tin_nhan.get("display_content", tin_nhan["content"])
            
            doan_van = tai_lieu.add_paragraph()
            doan_van.add_run(f"[{ten_vai_tro}]:\n").bold = True
            doan_van.add_run(f"{noi_dung}\n")
            
        luong_bo_nho = io.BytesIO()
        tai_lieu.save(luong_bo_nho)
        luong_bo_nho.seek(0)
        return luong_bo_nho

    # --- SUPABASE DATABASE ---
    def tai_danh_sach_chat_nguoi_dung(self, id_nguoi_dung):
        if not self.ket_noi_supabase:
            return {}
        try:
            ket_qua = self.ket_noi_supabase.table("chats").select("*").eq("user_id", str(id_nguoi_dung)).order("created_at", desc=False).execute()
            du_lieu_chats = {}
            for hang in ket_qua.data:
                id_chat = hang["id"]
                tieu_de = hang["title"]
                ket_qua_tin_nhan = self.ket_noi_supabase.table("messages").select("*").eq("chat_id", id_chat).order("created_at", desc=False).execute()
                du_lieu_chats[tieu_de] = {
                    "id": id_chat,
                    "messages": [
                        {
                            "role": m["role"],
                            "content": m["content"],
                            "display_content": m.get("display_content") or m["content"]
                        } for m in ket_qua_tin_nhan.data
                    ]
                }
            return du_lieu_chats
        except Exception as e:
            print(f"Lỗi nạp DB: {e}")
            return {}

    def luu_phien_chat(self, id_chat, tieu_de, id_nguoi_dung):
        if self.ket_noi_supabase:
            try:
                self.ket_noi_supabase.table("chats").upsert({"id": id_chat, "title": tieu_de, "user_id": str(id_nguoi_dung)}).execute()
            except Exception as e:
                print(f"Lỗi lưu phiên chat: {e}")

    def luu_tin_nhan(self, id_chat, vai_tro, noi_dung, noi_dung_hien_thi=""):
        if self.ket_noi_supabase:
            try:
                self.ket_noi_supabase.table("messages").insert({
                    "chat_id": id_chat,
                    "role": vai_tro,
                    "content": noi_dung,
                    "display_content": noi_dung_hien_thi
                }).execute()
            except Exception as e:
                print(f"Lỗi lưu tin nhắn: {e}")

    def xoa_phien_chat(self, id_chat):
        if self.ket_noi_supabase:
            try:
                self.ket_noi_supabase.table("chats").delete().eq("id", id_chat).execute()
            except Exception as e:
                print(f"Lỗi xóa phiên chat: {e}")

    def cap_nhat_tieu_de_chat(self, id_chat, tieu_de_moi):
        if self.ket_noi_supabase:
            try:
                self.ket_noi_supabase.table("chats").update({"title": tieu_de_moi}).eq("id", id_chat).execute()
            except Exception as e:
                print(f"Lỗi cập nhật tiêu đề: {e}")

    # --- ĐỌC CÁC TỆP ĐÍNH KÈM ---
    def doc_file_pdf(self, luong_bytes_file):
        try:
            if hasattr(luong_bytes_file, 'seek'):
                luong_bytes_file.seek(0)
            doc_pdf = PdfReader(luong_bytes_file)
            cac_trang_van_ban = []
            for i, trang in enumerate(doc_pdf.pages):
                van_ban = trang.extract_text()
                if van_ban and van_ban.strip():
                    cac_trang_van_ban.append(f"[Trang {i+1}]: {van_ban.strip()}")
            van_ban_day_du = "\n".join(cac_trang_van_ban)
            return van_ban_day_du if van_ban_day_du.strip() else "PDF_IS_SCANNED_IMAGE"
        except Exception as e:
            return f"Lỗi đọc file PDF: {str(e)}"

    def doc_file_word(self, luong_bytes_file):
        try:
            if hasattr(luong_bytes_file, 'seek'):
                luong_bytes_file.seek(0)
            
            # Đọc bằng docx tiêu chuẩn
            try:
                doc_word = docx.Document(luong_bytes_file)
                danh_sach_doan = [doan.text.strip() for doan in doc_word.paragraphs if doan.text and doan.text.strip()]
                for bang in doc_word.tables:
                    for hang in bang.rows:
                        hang_text = " | ".join([o.text.strip() for o in hang.cells if o.text.strip()])
                        if hang_text:
                            danh_sach_doan.append(hang_text)
                van_ban = "\n".join(danh_sach_doan)
                if van_ban.strip():
                    return van_ban
            except Exception:
                pass

            # Dự phòng đọc file tạo từ WPS Office
            import docx2txt
            if hasattr(luong_bytes_file, 'seek'):
                luong_bytes_file.seek(0)
            text_wps = docx2txt.process(luong_bytes_file)
            return text_wps.strip() if text_wps and text_wps.strip() else "FILE_WORD_RONG"
        except Exception as e:
            return f"Lỗi đọc file Word: {str(e)}"

    def doc_file_excel(self, luong_bytes_file):
        try:
            if hasattr(luong_bytes_file, 'seek'):
                luong_bytes_file.seek(0)
            bang_excel = pd.read_excel(luong_bytes_file, sheet_name=None)
            van_ban_trich_xuat = []
            for ten_sheet, df in bang_excel.items():
                van_ban_trich_xuat.append(f"--- Sheet: {ten_sheet} ---")
                van_ban_trich_xuat.append(df.to_markdown(index=False))
            return "\n".join(van_ban_trich_xuat)
        except Exception as e:
            return f"Lỗi đọc file Excel: {str(e)}"

    def phan_tich_hinh_anh(self, luong_bytes_file, yeu_cau="Đọc và trích xuất toàn bộ chữ viết có trong hình này."):
        if not self.danh_sach_keys:
            return "Chưa cấu hình API Key."
        
        if hasattr(luong_bytes_file, 'seek'):
            luong_bytes_file.seek(0)
        anh_base64 = base64.b64encode(luong_bytes_file.getvalue()).decode('utf-8')
        
        for key in self.danh_sach_keys:
            try:
                client_tam = Groq(api_key=key)
                phan_hoi = client_tam.chat.completions.create(
                    model=self.mo_hinh_hinh_anh,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": yeu_cau},
                                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{anh_base64}"}}
                            ]
                        }
                    ],
                    temperature=0.2,
                    max_tokens=2048
                )
                return phan_hoi.choices[0].message.content.strip()
            except Exception:
                continue
        return "Lỗi Vision AI: Không thể xử lý hình ảnh qua các Key hiện tại."

    # --- CHAT STREAMING CÓ TÍNH NĂNG XOAY API KEY TỰ ĐỘNG ---
    def luong_phan_hoi_chat(self, lich_su_tin_nhan_sach, ngu_canh_web=""):
        if not self.danh_sach_keys:
            yield "Chưa cấu hình API Key Groq trong Secrets."
            return

        huong_dan_he_thong = (
            "Bạn là Trợ lý AI giáo dục thông minh của Nhóm NEXUS. "
            "Trả lời chính xác, khoa học bằng tiếng Việt. "
            "Nếu không đủ thông tin, hãy thẳng thắn trả lời 'Tôi không đủ thông tin chắc chắn về vấn đề này', không tự đoán. "
            "Dùng Markdown và LaTeX ($...$) cho công thức toán/hóa."
        )
        
        if ngu_canh_web and ngu_canh_web != "CHƯA_CÓ_KẾT_QUẢ":
            huong_dan_he_thong += f"\n\n[DỮ LIỆU TRA CỨU WEB REAL-TIME]:\n{ngu_canh_web}\n\nTổng hợp thông tin trên để trả lời người dùng."

        tin_nhan_he_thong = {"role": "system", "content": huong_dan_he_thong}
        danh_sach_tin_nhan_gui = [tin_nhan_he_thong] + lich_su_tin_nhan_sach

        loi_chi_tiet = ""
        # Vòng lặp xoay API Key tự động khi gặp lỗi rate limit (429)
        for idx, key in enumerate(self.danh_sach_keys):
            key_sach = str(key).strip()
            if not key_sach:
                continue
            try:
                client_tam = Groq(api_key=key_sach)
                luong_phan_hoi = client_tam.chat.completions.create(
                    messages=danh_sach_tin_nhan_gui,
                    model=self.mo_hinh_van_ban,
                    temperature=0.0,
                    max_tokens=4000,
                    stream=True
                )
                for mieng_chu in luong_phan_hoi:
                    if mieng_chu.choices[0].delta.content:
                        yield mieng_chu.choices[0].delta.content
                return
            except Exception as e:
                loi_chi_tiet = str(e)
                if "429" in loi_chi_tiet or "rate_limit_exceeded" in loi_chi_tiet:
                    print(f"API Key số {idx+1} hết hạn mức! Đang chuyển Key tiếp theo...")
                    continue
                else:
                    yield f"Lỗi kết nối AI: {loi_chi_tiet}"
                    return

        yield f"Tất cả các API Key dự phòng đều gặp sự cố! Chi tiết: {loi_chi_tiet if loi_chi_tiet else 'Kiểm tra lại cấu hình Keys.'}"
